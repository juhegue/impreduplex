# -*- coding: utf-8 -*-

from io import BytesIO
import os
import sys
import glob
import platform
import subprocess
import tempfile
import base64
from PIL import Image
from pdf2image import convert_from_path
from fpdf import FPDF
from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
if platform.system() == 'Windows':
    import win32print
    import win32event
    import win32process
    import win32con
    import win32com.client
    from win32com.shell.shell import ShellExecuteEx
    from win32com.shell import shellcon

__version__ = '0.0.1'
appName = 'impreduplex'
author = 'juhegue'
date = 'jue 14 abr 2022'

img_relleno = """
iVBORw0KGgoAAAANSUhEUgAAABAAAAAQCAYAAAAf8/9hAAAABmJLR0QA/wD/AP+gvaeTAAAACXBI
WXMAAAsTAAALEwEAmpwYAAAAB3RJTUUH3wYHFSo7+aaN1AAAADNJREFUOMtj/P//PwMlgImBQkCx
ASzInJ07dxKlyd3dfRB5YdSAYWEA44Blpt+/f1PHBQDWcw8IAHFJDwAAAABJRU5ErkJggg==
"""
img_relleno = BytesIO(base64.b64decode(img_relleno.replace('\n', '')))


def win_duplex(nom_imp, duplex=3):
    print(f'Duplex({duplex}):{nom_imp}')
    printdefaults = {'DesiredAccess': win32print.PRINTER_ACCESS_USE}
    handle = win32print.OpenPrinter(nom_imp, printdefaults)
    level = 2
    attributes = win32print.GetPrinter(handle, level)
    antes = attributes['pDevMode'].Duplex
    attributes['pDevMode'].Duplex = duplex  # 1=no flip, 2=flip up, 3=flip over
    try:
        win32print.SetPrinter(handle, level, attributes, 0)
    except:
        pass

    return antes


def win_imprime(docu, impresora, duplex):
    if impresora.lower().endswith('.pdf'):
        wdFormatPDF = 17
        # word = win32com.client.Dispatch('Word.Application')
        word = win32com.client.gencache.EnsureDispatch('Word.Application')
        # word.Visible = False
        w = word.Documents.Open(docu)
        w.SaveAs(impresora, FileFormat=wdFormatPDF)
        w.Close()
        word.Quit()
    else:
        duplex_ant = None
        if impresora == 'ver':
            procInfo = ShellExecuteEx(nShow=win32con.SW_HIDE,
                                      fMask=shellcon.SEE_MASK_NOCLOSEPROCESS,
                                      lpVerb='open',
                                      lpFile=docu,
                                      )
        elif impresora == 'defecto':
            if duplex:
                impresora = win32print.GetDefaultPrinter()
                duplex_ant = win_duplex(impresora)

            procInfo = ShellExecuteEx(nShow=win32con.SW_HIDE,
                                      fMask=shellcon.SEE_MASK_NOCLOSEPROCESS,
                                      lpVerb='printto',
                                      lpFile=docu,
                                      )
        else:
            if duplex:
                duplex_ant = win_duplex(impresora)

            procInfo = ShellExecuteEx(nShow=win32con.SW_HIDE,
                                      fMask=shellcon.SEE_MASK_NOCLOSEPROCESS,
                                      lpVerb='printto',
                                      lpFile=docu,
                                      lpParameters='"%s"' % impresora
                                      )
        procHandle = procInfo['hProcess']
        obj = win32event.WaitForSingleObject(procHandle, win32event.INFINITE)
        rc = win32process.GetExitCodeProcess(procHandle)
        if duplex_ant:
            win_duplex(impresora, duplex_ant)


def win_get_paginas(documento):
    # TODO:: Para corregir el error
    # win32com\client\CLSIDToClass.pyc: KeyError: '{00020970-0000-0000-C000-000000000046}'
    if win32com.client.gencache.is_readonly:
        win32com.client.gencache.is_readonly = False
        win32com.client.gencache.Rebuild()  # create gen_py folder if needed

    # word = win32com.client.Dispatch('Word.Application')
    word = win32com.client.gencache.EnsureDispatch('Word.Application')
    # word.Visible = False
    w = word.Documents.Open(documento)
    w.Repaginate()
    paginas = w.ComputeStatistics(2)
    w.Close()
    word.Quit()
    return paginas


def escala_imagen(imagen, width, height):
    try:
        im = Image.open(imagen)
        w, h = im.size
        w = height * w / h
        if w > width:
            height = height * width / w
        else:
            width = w
    except:
        pass

    return int(width), int(height)


class MiFPDF(FPDF):

    def image(self, stream, x, y, w, h):
        tmp_file = os.path.join(tempfile.gettempdir(), next(tempfile._get_candidate_names())) + '.png'
        Image.open(stream).save(tmp_file)
        super().image(tmp_file, x=x, y=y, w=w, h=h, type='png', link='')
        os.remove(tmp_file)

    def footer(self):
        self.set_y(-6)
        self.set_font('Arial', 'I', 6)
        self.cell(0, 10, f'Página {self.page_no()}' + '/{nb}', 0, 0, 'C')


def albaranes_pdf(pdf, albaranes_img, albaran, img_pag_ancho, img_pag_alto):
    pdf.add_page()
    posy = 0
    for alto in range(0, img_pag_alto):
        posx = 0
        for ancho in range(0, img_pag_ancho):
            if albaran < len(albaranes_img):
                img = albaranes_img[albaran]
                w, h = escala_imagen(img, pdf.fw / img_pag_ancho, pdf.fh / img_pag_alto)
                inc_y = (pdf.fh / img_pag_alto - h) / 2
                inc_x = (pdf.fw / img_pag_ancho - w) / 2
                pdf.image(img, posx + inc_x, posy + inc_y, w, h)
                posx += pdf.fw / img_pag_ancho
                albaran += 1
        posy += pdf.fh / img_pag_alto
    return albaran


def crea_pdf(facturas_img, albaranes_img, doc_destino, img_pag_ancho, img_pag_alto):
    pdf = MiFPDF()
    pdf.set_margins(0, 0, 0)
    pdf.alias_nb_pages()
    pdf.set_author(author)
    pdf.set_creator(appName)

    albaran = 0
    for factura in facturas_img:
        pdf.add_page()
        pdf.image(factura, 0, 0, pdf.fw, pdf.fh)
        if albaran < len(albaranes_img):
            albaran = albaranes_pdf(pdf, albaranes_img, albaran, img_pag_ancho, img_pag_alto)

    while albaran < len(albaranes_img):
        albaran = albaranes_pdf(pdf, albaranes_img, albaran, img_pag_ancho, img_pag_alto)

    # siempre pares
    if pdf.page_no() % 2 != 0:
        pdf.add_page()

    pdf.output(doc_destino)


def albaranes_docx(document, albaranes_img, albaran, img_pag_ancho, img_pag_alto, twidth, theight):
    twidth -= .5
    theight -= img_pag_alto / 2 if platform.system() == 'Windows' else img_pag_alto  # para que no haga saltos de página
    table = document.add_table(rows=img_pag_alto, cols=img_pag_ancho)
    table.allow_autofit = False
    table.style = None
    for alto in range(0, img_pag_alto):
        for ancho in range(0, img_pag_ancho):
            cell = table.rows[alto].cells[ancho]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.style = None
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            if albaran < len(albaranes_img):
                width = twidth / img_pag_ancho
                height = theight / img_pag_alto
                imagen = albaranes_img[albaran]
                w, h = escala_imagen(imagen, width, height)
                run.add_picture(img_relleno, width=Cm(width), height=Cm((height-h)/2))
                run.add_picture(imagen, width=Cm(w), height=Cm(h))
                run.add_picture(img_relleno, width=Cm(width), height=Cm((height-h))/2)
                albaran = albaran + 1
    return albaran


def crea_docx(facturas_img, albaranes_img, doc_destino, img_pag_ancho, img_pag_alto):
    document = Document('impreduplex.docx')
    width = document.sections[0].page_width.cm
    height = document.sections[0].page_height.cm
    top = document.sections[0].top_margin.cm
    bottom = document.sections[0].bottom_margin.cm
    left = document.sections[0].left_margin.cm
    right = document.sections[0].right_margin.cm
    twidth = width - left - right
    theight = height - top - bottom

    for paragraph in document.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    paginas = albaran = 0
    for factura in facturas_img:
        paginas += 1
        document.add_picture(factura, width=Cm(twidth), height=Cm(theight))
        if albaran < len(albaranes_img):
            albaran = albaranes_docx(document, albaranes_img, albaran, img_pag_ancho, img_pag_alto, twidth, theight)
            paginas += 1

    while albaran < len(albaranes_img):
        albaran = albaranes_docx(document, albaranes_img, albaran, img_pag_ancho, img_pag_alto, twidth, theight)
        paginas += 1

    # Las páginas deben ser siempre pares
    if platform.system() == 'Windows':
        # en windows se cuentan las páginas con el Word
        document.save(doc_destino)
        paginas = win_get_paginas(doc_destino)
        print(f'Páginas creadas: {paginas}')
        if paginas % 2 != 0:
            document = Document(doc_destino)
            document.add_page_break()
            document.save(doc_destino)
            paginas = win_get_paginas(doc_destino)
            print(f'Página añadida, total: {paginas}')
    else:
        if paginas % 2 != 0:
            document.add_page_break()
            # document.add_paragraph('.')
        document.save(doc_destino)


def main():
    args = sys.argv[1:]
    if len(args) < 5:
        print('Error.')
        sys.exit(-1)

    file_factu = args[0]
    path_alba = args[1]
    doc_destino = args[2]
    img_pag_ancho = int(args[3])
    img_pag_alto = int(args[4])
    # ver=Visualiza y para windows: defecto=Impresora defecto o nombre impresora
    impresora = args[5] if len(args) > 5 else None
    duplex = True if len(args) > 6 else False

    facturas_img = list()
    imagenes = convert_from_path(file_factu)
    for imagen in imagenes:
        imagefile = BytesIO()
        imagen.save(imagefile, format='PNG')
        facturas_img.append(imagefile)

    albaranes_img = list()
    for albaran in sorted(glob.glob(path_alba)):
        with open(albaran, 'rb') as f:
            data = f.read()

        if data[1:4] == b'PDF':
            imagenes = convert_from_path(albaran)
            for imagen in imagenes:
                imagefile = BytesIO()
                imagen.save(imagefile, format='PNG')
                albaranes_img.append(imagefile)
        else:
            albaranes_img.append(BytesIO(data))

    if doc_destino.lower().endswith('.pdf'):
        crea_pdf(facturas_img, albaranes_img, doc_destino, img_pag_ancho, img_pag_alto)
    else:
        crea_docx(facturas_img, albaranes_img, doc_destino, img_pag_ancho, img_pag_alto)

    if impresora:
        if platform.system() == 'Darwin':       # mac
            subprocess.call(('open', doc_destino))
        elif platform.system() == 'Windows':
            win_imprime(doc_destino, impresora, duplex)
        else:                                   # linux
            subprocess.call(('xdg-open', doc_destino))


if __name__ == '__main__':
    main()

